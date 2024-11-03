import os
from docx import Document
from llms.openai_wrapper import openai_llm as llm
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


llm_selection = os.environ.get('LLM_SELECTION', "qwen2.5-instruct")
def convert_to_markdown(doc):
    """Converts a Word document to Markdown."""
    markdown_content = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Ignore empty paragraphs
            level = get_heading_level(paragraph)
            if level is not None:
                markdown_content = f"{markdown_content}{'#' * level} {text}\n\n"
            else:
                markdown_content = f"{markdown_content}{text}\n\n"
    for table in doc.tables:
        for row in table.rows:
            markdown_row = "|"
            for cell in row.cells:
                markdown_row = f"{markdown_row} {cell.text.strip()} |"
            markdown_content = f"{markdown_content}\n{markdown_row}\n"
    return markdown_content


def get_heading_level(paragraph):
    """Determines the heading level of a paragraph."""
    style = paragraph.style.name.lower()
    if 'heading' in style:
        try:
            return int(style.split(' ')[1])
        except:
            return None
    return None


def convert_from_markdown(markdown):
    """Converts Markdown back to a Word document format using the template's styles."""
    new_doc = Document()  # Create an empty document
    new_doc.styles['Normal'].font.name = u'宋体'
    new_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    new_doc.styles['Normal'].font.size = Pt(12)
    new_doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    lines = markdown.split('\n')
    lines = [line.strip() for line in lines if line.strip()]

    # Process the first line separately to ensure it's treated as the main title
    if lines and lines[0]:
        Head = new_doc.add_heading(level=1)
        Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = Head.add_run(f"{lines[0].lstrip('#').strip()}\n")
        run.font.name = u'Cambria'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

    # Process the rest of the lines
    current_table = None
    for line in lines[1:]:
        if line.startswith('#'):
            text = line.lstrip('#')
            level = len(line) - len(text)
            Head = new_doc.add_heading(level=level)
            run = Head.add_run(text.strip())
            run.font.name = u'Cambria'
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif '|' in line and line.count('|') >= 2:
            row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if current_table is None:
                # Start a new table if one doesn't already exist
                col_count = len(row_cells)
                current_table = new_doc.add_table(rows=1, cols=col_count)
            row = current_table.add_row().cells
            for i, cell_text in enumerate(row_cells):
                row[i].text = cell_text
        else:
            if current_table is not None:
                # Finish the current table before adding a new paragraph
                current_table = None
            new_doc.add_paragraph(line.strip())

    return new_doc


async def modify_word_document(file_path, modification_opinion, out_file, logger) -> (int, str):
    """Reads a Word document, modifies its content based on the given opinion, and writes it back."""
    if not os.path.exists(file_path) or not file_path.endswith('.docx'):
        logger.debug(f"{file_path}不存在或文件格式错误")
        return -5, "仅能处理docx文件喔~"

    doc = Document(file_path)
    markdown_content = convert_to_markdown(doc)
    if not markdown_content:
        return 0, "无法解析word文档，请简化格式并保证使用docx。"

    system_prompt = (f"你是一名来自上海市静安区临汾路街道的优秀社区工作者，请根据如下修改意见对输入的Markdown文本进行修改:\n"
                     f"{modification_opinion}\n\n"
                     f"注意：请不要修改Markdown结构与格式，只修改文本内容。")
    messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": markdown_content}]
    modified_markdown = await llm(messages=messages, model=llm_selection, logger=logger)
    if not modified_markdown:
        return -11, "LLM暂不可用，请联系管理员[苦涩]"

    if len(modified_markdown) <= 38:
        return 0, modified_markdown

    try:
        new_doc = convert_from_markdown(modified_markdown)
    except Exception as e:
        logger.debug(f"转换Markdown为Word文档时出错: {e}")
        return 0, modified_markdown

    new_doc.save(out_file)

    return 21, ''
