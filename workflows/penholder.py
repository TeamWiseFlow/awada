import os
from docx import Document
from llms.openai_wrapper import openai_llm as llm
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from llms.qanything import search, check_health


retrieve_sys = "来自{org}的知识助理。结合给定的初稿和用户的扩写需求提炼出最合适的写作主题。"
retrieve_prompt = """初稿如下：
{text}

扩写需求如下：
{demand}

请输出提炼出的写作主题（只输出写作主题，不要输出其他内容）：
"""

rewrite_sys = "来自{org}的知识助理，擅长进行专业写作。根据用户的扩写需求，结合从知识库中召回的参考资料，对初稿进行扩写。"
rewrite_prompt = """初稿（Markdown格式）如下：
{text}

从知识库中找到的参考资料如下：
{reference}

扩写需求如下：
{demand}

要求：1、保持初稿的结构和提纲不变，如果一级标题为空，请拟订并补充；2、从{org}的角度书写；3、输出依然使用Markdown格式，直接输出扩写后的Markdown文本，不要输出其他内容。"""

modify_sys = "来自{org}的知识助理，擅长进行专业写作。根据用户的修改意见，对原稿进行修改。"
modify_prompt = """原稿（Markdown格式）如下：
{text}

修改意见如下：
{demand}

要求：1、保持原稿的结构和提纲不变，如果一级标题为空，请拟订并补充；2、从{org}的角度书写；3、输出依然使用Markdown格式，直接输出扩写后的Markdown文本，不要输出其他内容。"""

query_rewrite_sys = "对用户的输入进行改写，从中提炼出合适的写作主题"
query_rewrite_prompt = """用户输入如下：
{demand}

请输出提炼出的写作主题（只输出写作主题，不要输出其他内容）：
"""

# todo 这里未来要替换为“按模板写作”功能，暂时先暂时固定写死一个最基础的模板
event_plan_template = """# 
## 一、背景

## 二、目标

## 三、方案

"""


class Penholder:
    def __init__(self, org: str, llm_model: str, wf_dir: str, logger):
        self.llm_model = llm_model
        self.wf_dir = wf_dir
        self.logger = logger
        self.retrieve_sys = retrieve_sys.format(org=org)
        self.rewrite_sys = rewrite_sys.format(org=org)
        self.modify_sys = modify_sys.format(org=org)
        self.org = org
        self.logger.info("Penholder Ready")

    def _convert_to_markdown(self, doc):
        """Converts a Word document to Markdown."""
        markdown_content = ""
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Ignore empty paragraphs
                level = self._get_heading_level(paragraph)
                if level is not None:
                    markdown_content = f"{markdown_content}{'#' * level} {text}\n\n"
                else:
                    markdown_content = f"{markdown_content}{text}\n\n"
        for table in doc.tables:
            for row in table.rows:
                markdown_row = "|"
                for cell in row.cells:
                    markdown_row = f"{markdown_row} {cell.text.strip()} |"
                markdown_content = f"{markdown_content}\n{markdown_row}\n"
        return markdown_content

    def _get_heading_level(self, paragraph):
        """Determines the heading level of a paragraph."""
        style = paragraph.style.name.lower()
        if 'heading' in style:
            try:
                return int(style.split(' ')[1])
            except:
                return None
        return None

    def _convert_from_markdown(self, markdown):
        """Converts Markdown back to a Word document format using the template's styles."""
        new_doc = Document()  # Create an empty document
        new_doc.styles['Normal'].font.name = u'宋体'
        new_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        new_doc.styles['Normal'].font.size = Pt(12)
        new_doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        lines = markdown.split('\n')
        lines = [line.strip() for line in lines if line.strip()]

        # Process the first line separately to ensure it's treated as the main title
        if lines and lines[0]:
            Head = new_doc.add_heading(level=1)
            Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = Head.add_run(f"{lines[0].lstrip('#').strip()}\n")
            run.font.name = u'Cambria'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

        # Process the rest of the lines
        current_table = None
        for line in lines[1:]:
            if line.startswith('#'):
                text = line.lstrip('#')
                level = len(line) - len(text)
                Head = new_doc.add_heading(level=level)
                run = Head.add_run(text.strip())
                run.font.name = u'Cambria'
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif '|' in line and line.count('|') >= 2:
                row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if current_table is None:
                    # Start a new table if one doesn't already exist
                    col_count = len(row_cells)
                    current_table = new_doc.add_table(rows=1, cols=col_count)
                row = current_table.add_row().cells
                for i, cell_text in enumerate(row_cells):
                    row[i].text = cell_text
            else:
                if current_table is not None:
                    # Finish the current table before adding a new paragraph
                    current_table = None
                new_doc.add_paragraph(line.strip())

        return new_doc

    async def writing(self, opinion: str, out_file: str, kbs: list[str], file_path: str = None) -> (int, str):
        """
        Reads a Word document, modifies its content based on the given opinion, and writes it back.
        file_path 对应原始文件，如果为 None，则直接按 opinion 判断调用何种模板 （todo 目前调用模板是写死只有一种模板的，后续需要增加使用llm判断调用何种模板）
        # 2024-11-07 新增写作前先从知识库中检索参考文本，并将其作为写作的参考
        """
        
        if file_path:
            if not os.path.exists(file_path) or not file_path.endswith('.docx'):
                self.logger.debug(f"{file_path}不存在或文件格式错误")
                return -1, "目前扩写技能仅能处理docx文件喔[尴尬]"

            doc = Document(file_path)
            markdown_content = self._convert_to_markdown(doc)
            if not markdown_content:
                return 2, "无法解析word文档，请重新上传并保证使用正确的docx格式（或简化初始文档）。"

            prompt = retrieve_prompt.format(text=markdown_content, demand=opinion)
            query = await llm(messages=[{"role": "system", "content": self.retrieve_sys},
                                        {"role": "user", "content": prompt}], model=self.llm_model, logger=self.logger)
        else:
            prompt = query_rewrite_prompt.format(demand=opinion)
            query = await llm(messages=[{"role": "system", "content": query_rewrite_sys},
                                        {"role": "user", "content": prompt}], model=self.llm_model, logger=self.logger)
            markdown_content = event_plan_template

        query = query.strip()
        if not query:
            return -11, "LLM暂不可用，请联系管理员[苦涩]"
        
        if not kbs:
            self.logger.warning("no knowledge base, so no search result")
            search_result = []
        else:
            flag, msg, search_result = await search(query, kb_ids=kbs, model=self.llm_model, wf_dir=self.wf_dir, logger=self.logger)
            if flag == -4:
                self.logger.error(f"Qanything out of service, code:{flag}, msg:{msg}")
                search_result = []
            elif flag > 200:
                self.logger.warning(f"Qanything search failed, code:{flag}, msg:{msg}")
                if check_health():
                    self.logger.info("health check ok, try again")
                    flag, msg, search_result = await search(query, kb_ids=kbs, model=self.llm_model, wf_dir=self.wf_dir, logger=self.logger)
                    self.logger.info(f"retry result:{flag}, {msg}")
                else:
                    self.logger.error("health check failed")
                    search_result = []

        documents = []
        references = set()
        for doc in search_result:
            if doc['type'] == 'faq':
                documents.append(f"<document>{doc['content']}</document>")
                references.add(f"FAQ_{doc['title']}")
            elif doc['type'] == 'wiseflow':
                documents.append(f"<document>{doc['content']}</document>")
                if doc['title'] or doc['source']:
                    references.add(f"{doc['title']}\n{doc['source']}")
            elif doc['type'] == 'file':
                documents.append(f"<document>{doc['content']}</document>")
                references.add(doc['title'])

        # 据说模型对首尾信息更敏感，所以要把得分高的放两端，姑且信之
        n = len(documents)
        doc_list = [None] * n
        for i in range(n):
            if i % 2 == 0:
                doc_list[i // 2] = documents[i]
            else:
                doc_list[-(i // 2 + 1)] = documents[i]

        rag_text = "\n\n".join(doc_list)
        self.logger.debug(rag_text)
        prompt = rewrite_prompt.format(text=markdown_content, reference=rag_text, demand=opinion, org=self.org)

        modified_markdown = await llm(messages=[{"role": "system", "content": self.rewrite_sys},
                                                {"role": "user", "content": prompt}], model=self.llm_model, logger=self.logger)
        if not modified_markdown:
            return -11, "LLM暂不可用，请联系管理员[苦涩]"

        if len(modified_markdown) <= 38:
            self.logger.warning("!!!rewrite process not effective, return original markdown!!!")
            return 0, modified_markdown

        if references:
            modified_markdown = f"{modified_markdown}\n##参考资料：\n"
            for i, ref in enumerate(references):
                modified_markdown = f"{modified_markdown}{i+1}. {ref}\n"
        else:
            if not documents:
                modified_markdown = f"{modified_markdown}\n###未找到任何参考资料，请谨慎使用本生成文档！"

        try:
            new_doc = self._convert_from_markdown(modified_markdown)
        except Exception as e:
            self.logger.error(f"转换Markdown为Word文档时出错: {e}")
            return 0, modified_markdown

        new_doc.save(out_file)
        return 21, modified_markdown

    async def modify(self, opinion: str, out_file: str, history: str) -> (int, str):
        prompt = modify_prompt.format(text=history, demand=opinion, org=self.org)
        modified_markdown = await llm(messages=[{"role": "system", "content": self.modify_sys},
                                                {"role": "user", "content": prompt}], model=self.llm_model, logger=self.logger)
        if not modified_markdown:
            return -11, "LLM暂不可用，请联系管理员[苦涩]"

        try:
            new_doc = self._convert_from_markdown(modified_markdown)
        except Exception as e:
            self.logger.error(f"转换Markdown为Word文档时出错: {e}")
            return 0, modified_markdown

        new_doc.save(out_file)
        return 21, modified_markdown
